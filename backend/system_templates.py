"""System-seeded DOCX templates for each industry/subdomain.

These are built once on first server start and made available read-only to
every user via /api/templates (with is_system=true flag). Users can also
upload their own templates as before.

Templates are generated in-memory using python-docx with Romanian gas
engineering boilerplate text + placeholders. Replace these later with
your real legal forms.
"""
import io
import base64
from datetime import datetime, timezone
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def _h(doc, text, size=18):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(size)


def _p(doc, text, bold=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(11)
    if bold:
        r.bold = True


def _build_cerere_racordare() -> bytes:
    d = Document()
    _h(d, "CERERE DE RACORDARE LA REȚEAUA DE DISTRIBUȚIE GAZE NATURALE")
    d.add_paragraph()
    _p(d, "Către: {{osd}}")
    d.add_paragraph()
    _p(d, "Subsemnatul/Subscrisa {{beneficiar}}, cu domiciliul/sediul în localitatea {{localitate}}, județul {{judet}}, adresa {{adresa_lucrare}}, telefon {{telefon}}, email {{email}}, vă rog să aprobați racordarea la rețeaua de distribuție gaze naturale a imobilului situat la adresa menționată.")
    d.add_paragraph()
    _p(d, "Date tehnice ale lucrării:", bold=True)
    _p(d, "• Tip lucrare: {{tip_lucrare}}")
    _p(d, "• Debit instalat: {{debit_instalat}} mc/h")
    _p(d, "• Putere instalată: {{putere_instalata_kw}} kW")
    _p(d, "• Lungime estimată branșament: {{lungime_bransament}} m")
    _p(d, "• Presiune regim solicitată: {{presiune_regim}}")
    _p(d, "• Contor propus: {{contor_orientativ}}")
    d.add_paragraph()
    _p(d, "Date contract: nr. {{numar_contract}} din {{data_contract}}.")
    _p(d, "Proiectant: {{proiectant}}")
    _p(d, "Executant: {{executant}}")
    d.add_paragraph()
    _p(d, "Observații: {{observatii}}")
    d.add_paragraph()
    _p(d, "Data: {{data_document}}")
    _p(d, "Semnătura beneficiar: ________________________")
    out = io.BytesIO()
    d.save(out)
    return out.getvalue()


def _build_memoriu_tehnic() -> bytes:
    d = Document()
    _h(d, "MEMORIU TEHNIC — INSTALAȚIE GAZE NATURALE", size=16)
    d.add_paragraph()
    _p(d, "1. DATE GENERALE", bold=True)
    _p(d, "Beneficiar: {{beneficiar}}")
    _p(d, "Adresa lucrării: {{adresa_lucrare}}, {{localitate}}, jud. {{judet}}")
    _p(d, "Telefon contact: {{telefon}} | Email: {{email}}")
    _p(d, "Operator distribuție (OSD): {{osd}}")
    _p(d, "Tip lucrare: {{tip_lucrare}}")
    _p(d, "Contract: nr. {{numar_contract}} / {{data_contract}}")
    d.add_paragraph()
    _p(d, "2. ECHIPA TEHNICĂ AUTORIZATĂ", bold=True)
    _p(d, "Proiectant ANRE: {{proiectant}}")
    _p(d, "Executant ANRE: {{executant}}")
    _p(d, "Verificator documentație (VGD): {{verificator_vgd}}")
    _p(d, "Responsabil tehnic execuție (RTE): {{responsabil_rte}}")
    d.add_paragraph()
    _p(d, "3. DATE TEHNICE", bold=True)
    _p(d, "Debit instalat: {{debit_instalat}} mc/h")
    _p(d, "Debit calculat: {{debit_calculat_mc_h}} mc/h")
    _p(d, "Debit recomandat (cu marjă 10%): {{debit_recomandat_mc_h}} mc/h")
    _p(d, "Putere instalată estimată: {{putere_instalata_kw}} kW")
    _p(d, "Presiune regim: {{presiune_regim}}")
    _p(d, "Diametru conductă: {{diametru_conducta}}")
    _p(d, "Material conductă: {{material_conducta}}")
    _p(d, "Lungime branșament: {{lungime_bransament}} m")
    _p(d, "Punct racordare: {{punct_racordare}}")
    _p(d, "Post reglare: {{post_reglare}}")
    _p(d, "Contor recomandat: {{contor_orientativ}}")
    _p(d, "Risc presiune: {{risc_presiune}}")
    _p(d, "Estimare cost branșament: {{estimare_cost}} RON")
    d.add_paragraph()
    _p(d, "4. OBSERVAȚII TEHNICE", bold=True)
    _p(d, "{{observatii}}")
    d.add_paragraph()
    _p(d, "Data emiterii: {{data_document}}")
    _p(d, "Întocmit: {{proiectant}}")
    out = io.BytesIO()
    d.save(out)
    return out.getvalue()


def _build_borderou() -> bytes:
    d = Document()
    _h(d, "BORDEROU DOCUMENTE — DOSAR TEHNIC")
    d.add_paragraph()
    _p(d, "Beneficiar: {{beneficiar}}")
    _p(d, "Adresa lucrare: {{adresa_lucrare}}, {{localitate}}, {{judet}}")
    _p(d, "Tip lucrare: {{tip_lucrare}}")
    _p(d, "Contract: {{numar_contract}} / {{data_contract}}")
    d.add_paragraph()
    _p(d, "Conținut dosar:", bold=True)
    _p(d, "1. Cerere de racordare")
    _p(d, "2. Memoriu tehnic")
    _p(d, "3. Plan de încadrare (anexă)")
    _p(d, "4. Plan de situație (anexă)")
    _p(d, "5. Schemă izometrică (anexă)")
    _p(d, "6. Acord acces / declarație proprietate (anexă)")
    _p(d, "7. Verificare VGD: {{verificator_vgd}}")
    _p(d, "8. Confirmare RTE: {{responsabil_rte}}")
    d.add_paragraph()
    _p(d, "Întocmit: {{proiectant}} | Data: {{data_document}}")
    out = io.BytesIO()
    d.save(out)
    return out.getvalue()


def _build_adresa_osd() -> bytes:
    d = Document()
    _h(d, "ADRESĂ CĂTRE OPERATORUL SISTEMULUI DE DISTRIBUȚIE", size=14)
    d.add_paragraph()
    _p(d, "Către: {{osd}}")
    _p(d, "În atenția: Departamentul Racordări")
    d.add_paragraph()
    _p(d, "Stimată/Stimate domnule director,")
    d.add_paragraph()
    _p(d, "Prin prezenta, vă transmitem documentația tehnică aferentă lucrării de {{tip_lucrare}} pentru beneficiarul {{beneficiar}}, situat la adresa {{adresa_lucrare}}, {{localitate}}, județul {{judet}}.")
    d.add_paragraph()
    _p(d, "Detalii lucrare:", bold=True)
    _p(d, "• Contract nr. {{numar_contract}} din {{data_contract}}")
    _p(d, "• Debit instalat: {{debit_instalat}} mc/h")
    _p(d, "• Lungime branșament: {{lungime_bransament}} m")
    _p(d, "• Proiectant ANRE: {{proiectant}}")
    _p(d, "• Executant ANRE: {{executant}}")
    d.add_paragraph()
    _p(d, "Vă rugăm să analizați documentația și să comunicați avizul dvs. în termenul legal.")
    d.add_paragraph()
    _p(d, "Cu deosebită stimă,")
    _p(d, "{{proiectant}}")
    _p(d, "Data: {{data_document}}")
    out = io.BytesIO()
    d.save(out)
    return out.getvalue()


def _build_certificare_vgd() -> bytes:
    d = Document()
    _h(d, "CERTIFICARE INTERNĂ — VERIFICATOR DOCUMENTAȚIE (VGD)", size=14)
    d.add_paragraph()
    _p(d, "Beneficiar lucrare: {{beneficiar}}")
    _p(d, "Adresa lucrării: {{adresa_lucrare}}, {{localitate}}, {{judet}}")
    _p(d, "Tip lucrare: {{tip_lucrare}}")
    _p(d, "Contract: nr. {{numar_contract}} / {{data_contract}}")
    d.add_paragraph()
    _p(d, "1. VERIFICATOR DOCUMENTAȚIE", bold=True)
    _p(d, "Nume: {{verificator_vgd}}")
    _p(d, "Atestat ANRE: {{atestat_vgd}}")
    _p(d, "Data verificării: {{data_verificare_vgd}}")
    _p(d, "Status verificare: {{status_vgd}}")
    d.add_paragraph()
    _p(d, "2. OBSERVAȚII VERIFICARE", bold=True)
    _p(d, "{{observatii_vgd}}")
    d.add_paragraph()
    _p(d, "Subsemnatul, în calitate de Verificator Documentație autorizat ANRE, certific verificarea completă a documentației tehnice aferente lucrării de mai sus, conform reglementărilor în vigoare.")
    d.add_paragraph()
    _p(d, "Data certificării: {{data_document}}")
    _p(d, "Semnătură VGD: ________________________")
    _p(d, "Ștampilă: ________________________")
    out = io.BytesIO()
    d.save(out)
    return out.getvalue()


def _build_certificare_rte() -> bytes:
    d = Document()
    _h(d, "CERTIFICARE INTERNĂ — RESPONSABIL TEHNIC EXECUȚIE (RTE)", size=14)
    d.add_paragraph()
    _p(d, "Beneficiar lucrare: {{beneficiar}}")
    _p(d, "Adresa lucrării: {{adresa_lucrare}}, {{localitate}}, {{judet}}")
    _p(d, "Tip lucrare: {{tip_lucrare}}")
    _p(d, "Executant: {{executant}}")
    _p(d, "Contract: nr. {{numar_contract}} / {{data_contract}}")
    d.add_paragraph()
    _p(d, "1. RESPONSABIL TEHNIC EXECUȚIE", bold=True)
    _p(d, "Nume: {{responsabil_rte}}")
    _p(d, "Autorizație ANRE: {{autorizatie_rte}}")
    _p(d, "Data verificării execuției: {{data_verificare_rte}}")
    _p(d, "Status execuție: {{status_rte}}")
    d.add_paragraph()
    _p(d, "2. OBSERVAȚII EXECUȚIE", bold=True)
    _p(d, "{{observatii_rte}}")
    d.add_paragraph()
    _p(d, "Subsemnatul, în calitate de Responsabil Tehnic cu Execuția autorizat ANRE, certific execuția conformă a lucrării de mai sus și respectarea documentației tehnice verificate.")
    d.add_paragraph()
    _p(d, "Data certificării: {{data_document}}")
    _p(d, "Semnătură RTE: ________________________")
    _p(d, "Ștampilă: ________________________")
    out = io.BytesIO()
    d.save(out)
    return out.getvalue()


# ====================== FOTOVOLTAICE TEMPLATES ======================

def _build_avc_prosumator_fv() -> bytes:
    """Aviz tehnic de racordare pentru prosumator FV (< 27 kWp casnic)."""
    d = Document()
    _h(d, "CERERE AVIZ TEHNIC DE RACORDARE — PROSUMATOR FOTOVOLTAIC")
    d.add_paragraph()
    _p(d, "Către: {{ose}}")
    d.add_paragraph()
    _p(d, "Subsemnatul/Subscrisa {{beneficiar}}, cu domiciliul/sediul în {{adresa_lucrare}}, localitatea {{localitate}}, județul {{judet}}, telefon {{telefon}}, email {{email}}, vă rog să eliberați Aviz Tehnic de Racordare (ATR) pentru instalația fotovoltaică proprie.")
    d.add_paragraph()
    _p(d, "Date tehnice ale instalației FV:", bold=True)
    _p(d, "• Putere instalată DC: {{putere_dc_kwp}} kWp")
    _p(d, "• Putere maximă AC injectată: {{putere_ac_kw}} kW")
    _p(d, "• Număr panouri: {{numar_panouri}}")
    _p(d, "• Putere unitară panou: {{putere_panou_wp}} Wp")
    _p(d, "• Marcă invertor: {{marca_invertor}}")
    _p(d, "• Producție estimată anuală: {{productie_anuala_kwh}} kWh/an")
    _p(d, "• Tip racordare: {{tip_racordare}} (monofazat / trifazat)")
    _p(d, "• Sistem stocare BESS: {{stocare_kwh}} kWh (dacă există)")
    d.add_paragraph()
    _p(d, "Date proiectant:", bold=True)
    _p(d, "Nume societate: {{nume_societate}}")
    _p(d, "CUI: {{cui_societate}}")
    _p(d, "Atestat ANRE: {{atestat_anre}}")
    _p(d, "Reprezentant legal: {{reprezentant_legal}}")
    d.add_paragraph()
    _p(d, "Anexe:", bold=True)
    _p(d, "• Schema electrică monofilară")
    _p(d, "• Certificat de urbanism")
    _p(d, "• Extras de carte funciară")
    _p(d, "• Memoriu tehnic + dimensionare")
    _p(d, "• Buletin energetic (dacă e cazul)")
    d.add_paragraph()
    _p(d, "Data: {{data_document}}")
    _p(d, "Semnătura solicitant: ____________________")
    out = io.BytesIO()
    d.save(out)
    return out.getvalue()


def _build_memoriu_fv_rezidential() -> bytes:
    d = Document()
    _h(d, "MEMORIU TEHNIC — INSTALAȚIE FOTOVOLTAICĂ REZIDENȚIALĂ")
    d.add_paragraph()
    _p(d, "1. DATE GENERALE", bold=True)
    _p(d, "Denumirea lucrării: Instalație FV rezidențială pentru autoconsum + injecție (prosumator)")
    _p(d, "Amplasament: {{adresa_lucrare}}, localitatea {{localitate}}, județul {{judet}}")
    _p(d, "Beneficiar: {{beneficiar}}")
    _p(d, "Tip clădire: {{tip_cladire}}")
    _p(d, "Suprafață disponibilă acoperiș: {{suprafata_acoperis_mp}} mp")
    _p(d, "Orientare azimut: {{azimut}}°")
    _p(d, "Înclinație panouri: {{inclinatie}}°")
    d.add_paragraph()
    _p(d, "2. DESCRIEREA INSTALAȚIEI", bold=True)
    _p(d, "Putere instalată DC: {{putere_dc_kwp}} kWp")
    _p(d, "Putere maximă AC: {{putere_ac_kw}} kW")
    _p(d, "Tehnologie panouri: {{tehnologie_panouri}}")
    _p(d, "Tip invertor: {{tip_invertor}}")
    _p(d, "Eficiență invertor: {{randament_invertor}} %")
    d.add_paragraph()
    _p(d, "3. DIMENSIONARE & RANDAMENT", bold=True)
    _p(d, "Iradiere medie anuală: {{iradiere_kwh_mp_an}} kWh/mp/an")
    _p(d, "Producție specifică: {{productie_specifica}} kWh/kWp/an")
    _p(d, "Producție anuală estimată: {{productie_anuala_kwh}} kWh")
    _p(d, "Consum anual locuință: {{consum_anual_kwh}} kWh")
    _p(d, "Grad autoconsum estimat: {{grad_autoconsum}} %")
    d.add_paragraph()
    _p(d, "4. PROTECȚII & SIGURANȚĂ", bold=True)
    _p(d, "Protecții AC: întreruptor diferențial 30 mA + supratensiuni clasă II")
    _p(d, "Protecții DC: siguranțe DC, paratrăsnet conform IEC 62305")
    _p(d, "Anti-islanding: conform NTS/2021 ANRE")
    d.add_paragraph()
    _p(d, "5. CONFORMITATE LEGALĂ", bold=True)
    _p(d, "OUG 163/2022 — Programul Casa Verde Fotovoltaic")
    _p(d, "Ord. ANRE 228/2018 — racordare prosumator")
    _p(d, "Ord. ANRE 102/2023 — modificări recente")
    _p(d, "PE 124/95 — proiectare instalații electrice")
    d.add_paragraph()
    _p(d, "Proiectant: {{nume_proiectant}}, atestat ANRE: {{atestat_anre}}")
    _p(d, "Data: {{data_document}}")
    out = io.BytesIO()
    d.save(out)
    return out.getvalue()


def _build_borderou_fv() -> bytes:
    d = Document()
    _h(d, "BORDEROU DOCUMENTE DOSAR PROSUMATOR FOTOVOLTAIC")
    d.add_paragraph()
    _p(d, "Solicitant: {{beneficiar}}")
    _p(d, "Adresă: {{adresa_lucrare}}, {{localitate}}, {{judet}}")
    d.add_paragraph()
    _p(d, "Documente componente dosar:", bold=True)
    for line in [
        "1. Cerere ATR (Aviz Tehnic de Racordare)",
        "2. Memoriu tehnic instalație FV",
        "3. Schema electrică monofilară",
        "4. Schema unifilară interior",
        "5. Plan situație + plan amplasament",
        "6. Certificat de urbanism",
        "7. Extras de carte funciară (max. 30 zile)",
        "8. Acord vecini (dacă structura impune)",
        "9. Buletin energetic clădire",
        "10. Schema protecții AC/DC",
        "11. Fișa tehnică panouri",
        "12. Fișa tehnică invertor",
        "13. Certificat garanție echipamente",
        "14. Atestat ANRE proiectant",
        "15. Atestat ISCIR montator (dacă > 10 kWp)",
    ]:
        _p(d, line)
    d.add_paragraph()
    _p(d, "Proiectant: {{nume_proiectant}}")
    _p(d, "Data: {{date_proiect_data}}")
    out = io.BytesIO()
    d.save(out)
    return out.getvalue()


SYSTEM_TEMPLATES = [
    {
        "key": "sys_cerere_racordare_gaz",
        "name": "Cerere racordare gaze naturale (sistem)",
        "industry": "gas_engineering",
        "subdomain": "bransamente_gaz",
        "builder": _build_cerere_racordare,
    },
    {
        "key": "sys_memoriu_tehnic_gaz",
        "name": "Memoriu tehnic instalație gaze (sistem)",
        "industry": "gas_engineering",
        "subdomain": "instalatii_utilizare",
        "builder": _build_memoriu_tehnic,
    },
    {
        "key": "sys_borderou_documente_gaz",
        "name": "Borderou documente dosar gaze (sistem)",
        "industry": "gas_engineering",
        "subdomain": "bransamente_gaz",
        "builder": _build_borderou,
    },
    {
        "key": "sys_adresa_osd_gaz",
        "name": "Adresă către OSD (sistem)",
        "industry": "gas_engineering",
        "subdomain": "bransamente_gaz",
        "builder": _build_adresa_osd,
    },
    {
        "key": "sys_certificare_vgd_gaz",
        "name": "Certificare internă VGD (sistem)",
        "industry": "gas_engineering",
        "subdomain": "bransamente_gaz",
        "builder": _build_certificare_vgd,
    },
    {
        "key": "sys_certificare_rte_gaz",
        "name": "Certificare internă RTE (sistem)",
        "industry": "gas_engineering",
        "subdomain": "bransamente_gaz",
        "builder": _build_certificare_rte,
    },
    {
        "key": "sys_atr_prosumator_fv",
        "name": "Cerere ATR prosumator FV (sistem)",
        "industry": "photovoltaic",
        "subdomain": "instalatii_fv_rezidential",
        "builder": _build_avc_prosumator_fv,
    },
    {
        "key": "sys_memoriu_fv_rezidential",
        "name": "Memoriu tehnic FV rezidențial (sistem)",
        "industry": "photovoltaic",
        "subdomain": "instalatii_fv_rezidential",
        "builder": _build_memoriu_fv_rezidential,
    },
    {
        "key": "sys_borderou_fv",
        "name": "Borderou dosar prosumator FV (sistem)",
        "industry": "photovoltaic",
        "subdomain": "instalatii_fv_rezidential",
        "builder": _build_borderou_fv,
    },
]


async def seed_system_templates(db):
    """Ensure all system templates exist (idempotent)."""
    now = datetime.now(timezone.utc).isoformat()
    from docx_processor import extract_placeholders
    for tpl in SYSTEM_TEMPLATES:
        existing = await db.system_templates.find_one({"key": tpl["key"]})
        data = tpl["builder"]()
        placeholders = extract_placeholders(data)
        doc = {
            "key": tpl["key"],
            "template_id": tpl["key"],
            "name": tpl["name"],
            "industry": tpl["industry"],
            "subdomain": tpl["subdomain"],
            "placeholders": placeholders,
            "size_bytes": len(data),
            "data_b64": base64.b64encode(data).decode("ascii"),
            "is_system": True,
            "created_at": existing["created_at"] if existing else now,
            "updated_at": now,
        }
        if existing:
            await db.system_templates.update_one({"key": tpl["key"]}, {"$set": doc})
        else:
            await db.system_templates.insert_one(doc)
